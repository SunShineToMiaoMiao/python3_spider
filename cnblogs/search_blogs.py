import logging;logging.basicConfig(level=logging.INFO)
import urllib.parse
import urllib.request
import re
from bs4 import BeautifulSoup

import os, datetime, time, json
from docx import Document
from docx.shared import Pt, RGBColor, Inches
# from docx.enum.text import WD_BREAK_TYPE, WD_BREAK, WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

import keyword


'''
保存搜索到的每一篇博客 

功能(按照正文的每一段来抓取)：

    p标签的文字内容
    图片
    保存为word文档
    序列元素
    将代码段保存为对应样式

缺点：
    无法爬取结构繁多的div标签
    复用性太差

version 1.0.0
'''
# region code

# 获取python的关键字
python_keyword = keyword.kwlist

# 关键字匹配的正则
keyword_re = '(' + '|'.join(python_keyword) + ')'


# '----根据代码内容设置对应的颜色----'
# pre标签|code标签内没有span标签，仅是str
def set_code_color(pre_tag, paragraph):
    span_con = pre_tag.text.split('\n')
    for c in span_con:
        run = paragraph.add_run(c)
        run.font.size = Pt(12)
        run.font.name = 'Consolas'

        if re.match('(# |-- |// ).+', c):  # 注释
            run.font.color.rgb = RGBColor(0, 128, 0)
        elif re.match(keyword_re, c):  # 关键词
            run.font.color.rgb = RGBColor(0, 0, 255)
        run.add_break(6)
    return paragraph


# '----根据代码内容设置对应的颜色----'
# pre标签内全是span标签
def set_code_color_of_span(span_tags, paragraph):
    for c in span_tags:
        # 若是pre内包含了一个外层span标签，去掉
        par_span = c.find_all(lambda x: x.name != '', recursive=False)
        if len(par_span) > 1:
            continue

        run = paragraph.add_run(c.text)
        run.font.size = Pt(12)
        run.font.name = 'Consolas'

        # 根据span的class或style来分离颜色
        if c.has_attr('class'):
            if c.attrs['class'] == 'hljs-comment':
                run.font.color.rgb = RGBColor(0, 128, 0)
            elif c.attrs['class'] == 'hljs-keyword':
                run.font.color.rgb = RGBColor(0, 0, 255)
        elif c.has_attr('style'):
            span_style = re.match('color: #.{6}', c.attrs['style'])
            if span_style and span_style.group()[7:] == '#0000ff':
                run.font.color.rgb = RGBColor(0, 0, 255)
            elif span_style and span_style.group()[7:] == '#008000':
                run.font.color.rgb = RGBColor(0, 128, 0)
            elif span_style and span_style.group()[7:] == '#800000':
                run.font.color.rgb = RGBColor(128, 0, 0)
    return paragraph


# '处理pre标签--每篇博客内的代码部分'
def get_docx_of_pre_tag(docx, pre_tag):
    # code标签
    code_tag = pre_tag.find('code')
    # span标签
    span_tags = pre_tag.find_all('span')
    # 新建段落
    par = docx.add_paragraph()

    if code_tag:  # code标签
        set_code_color(code_tag, par)
    elif len(span_tags) == 0:  # 没有span的pre标签
        set_code_color(pre_tag, par)
    else:  # span标签
        set_code_color_of_span(span_tags, par)

    return docx


# 获取ul标签的内容
def get_ul_content(ul_tag, docx):
    li_string = ''

    li_tags = ul_tag.find_all('li')
    if len(li_tags) != 0:
        for c in li_tags:
            # 增加无序列表
            docx.add_paragraph(c.text, style='List Bullet')
    return li_string


# '创建目录'
def create_dir():
    dir = datetime.datetime.now().strftime('%Y%m%d%H%M')
    path = os.path.join('doc', dir)
    if os.path.exists(path):
        # os.remove(path)
        return path
    else:
        os.mkdir(path)
        return path


# '写入txt--str'
def save_str_txt(data, file_path):
    fp = open(file_path, 'w+', encoding='utf-8')
    try:
        fp.write(data)
    finally:
        fp.close()


# '写入txt--JSON'
def save_json_txt(list_name, file_path):
    # 一定要加上ensure_ascii=False, 才能保存为汉字
    data = json.dumps(list_name, ensure_ascii=False)
    with open(file_path, 'w+', encoding='utf-8') as f:
        f.write(data)


# '根据img_url保存图片文件'
def save_img_file(img_url, dir_path):
    mirS = str(datetime.datetime.now().microsecond)
    f_str = str(time.strftime('%H_%M_%S', time.localtime()) + '__' + mirS + '.png')
    file_path = os.path.join(dir_path, f_str)
    con = urllib.request.urlopen(img_url).read()
    with open(file_path, 'wb') as f:
        f.write(con)
    return file_path


# '保存为word文档'
def save_docx(url):
    # 新建word文档
    docx = Document()
    # 设置中文字体
    docx.styles['Normal'].font.name = '宋体'
    docx.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    # paragraph_format.line_spacing = Pt(18)

    # 保存目录
    dir_path = create_dir()

    # 获取html文档树
    blog_content = urllib.request.urlopen(url).read()
    soup = BeautifulSoup(blog_content, 'html.parser')
    post = soup.find('div', attrs={'class': 'post'})

    if post:

        #  标题
        title_a = post.find('a', attrs={'id': 'cb_post_title_url'})
        if title_a:
            title = title_a.string
            docx.add_heading(title, level=0)
            logging.info('已添加文档标题：%s' % title)

        docx.add_paragraph(url)
        logging.info('已添加文档链接：%s' % url)

        # 正文
        post_body = post.find('div', attrs={'id': 'cnblogs_post_body'})
        if post_body:
            # 判断是否有子节点
            body_child_tags = post_body.find_all(lambda x: x.name != '', recursive=False)
            # 遍历子节点
            for child in body_child_tags:
                # logging.info('正文每一小节html：%s' % child)

                if child.name == 'p':  # p标签
                    docx.add_paragraph(child.text)

                # 判断是否有子节点
                child_tag = child.find_all(lambda x: x.name != '', recursive=False)
                if len(child_tag) == 0:
                    if re.match('h[1-6]{1}', child.name):  # h1到h6标签
                        h_tag_level = child.name[-1]
                        docx.add_heading(child.string, level=int(h_tag_level))
                else:
                    if child.name == 'pre':  # 代码
                        get_docx_of_pre_tag(docx, child)
                    elif child.name == 'div' and child.attrs['class'][0] == 'cnblogs_code':  # div内的代码
                        pre_tag = child.find('pre')
                        get_docx_of_pre_tag(docx, pre_tag)
                    elif child.name == 'address':  # 地址
                        docx.add_paragraph(child.text)
                    elif child.name == 'ul':  # 无序列表
                        get_ul_content(child, docx)

                    # 图片
                    img_tag = child.find('img')
                    if img_tag:
                        img_name = save_img_file(img_tag.attrs['src'], dir_path)
                        docx.add_picture(img_name, width=Inches(6))

        # 保存文件
        docx_path = os.path.join(dir_path, title)
        docx.save(docx_path + '.docx')
        logging.info('已保存的word路径：%s' % docx_path + '.docx')


# 根据查询条件返回HTML
def request_html_by_query(url, **kwargs):
    # user_agent = 'Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/58.0.3029.81 Safari/537.36'
    # headers = {'User-Agent': user_agent}
    data = urllib.parse.urlencode(kwargs)
    api = url + '?' + data

    logging.info('访问链接：%s' % api)

    response_result = urllib.request.urlopen(api).read()
    html = response_result.decode('utf-8')
    return html


# '获取搜索到的博客数组'
def get_search_arr(html, num):
    soup = BeautifulSoup(html, 'html.parser')
    all_div = soup.find_all('div',attrs={'class': 'searchItem'}, limit=num)

    logging.info('共提取出%s个包含原文链接的div' % len(all_div))

    blogs = []
    for item in all_div:
        blog = get_search_item(item)
        blogs.append(blog)
    return blogs


# 按class查找单个标签
def s_class_find(item, tag, c):
    return item.find(tag, attrs={'class':c})


# '解析每个内容div'
def get_search_item(item):
    rs = {}

    a_title = s_class_find(item, 'h3', 'searchItemTitle')
    if a_title:
        # 博客标题
        rs['title'] = a_title.text
        rs['href'] = a_title.find('a')['href']

    # p_summary = s_class_find(item, 'span', 'searchCon')
    # if p_summary:
    #     rs['summary'] = p_summary.text

    userName = s_class_find(item, 'span', 'searchItemInfo-userName')
    rs['author'] = userName.text
    rs['author_url'] = userName.find('a')['href']

    published_time = s_class_find(item, 'span', 'searchItemInfo-publishDate')
    rs['published_time'] = published_time.text

    view_num = s_class_find(item, 'span', 'searchItemInfo-views')
    rs['view_num'] = re.search(r'\d+', view_num.text).group()

    comment_num = s_class_find(item, 'span', 'searchItemInfo-comments')
    if comment_num:
        rs['comment_num'] = re.search(r'\d+', comment_num.text).group()

    recommend_num = s_class_find(item, 'span', 'searchItemInfo-good')
    if recommend_num:
        rs['recommend_num'] = re.search(r'\d+', recommend_num.text).group()

    logging.info('解析到的单个div：%s' % rs)
    return rs


# 保存搜索到的每一篇博客
def save_blogs_by_one(blogs, num):
    logging.info('数据列表：%s' % blogs)

    # 保存搜索到的博客目录
    save_json_txt(blogs, os.path.join('doc', 'blogs_directory'+'_'+ num +'.txt'))

    for i in blogs:
        href = i['href']
        # 保存为word文档
        save_docx(href)

# endregion


if __name__ == '__main__':
    # http://zzk.cnblogs.com/s/blogpost?Keywords=python%20%E7%88%AC%E8%99%AB
    # http://zzk.cnblogs.com/s/blogpost?Keywords= &pageindex=1

    # 下载单个blog
    # save_docx('http://www.cnblogs.com/shiyanlou/p/7985806.html')

    url = 'http://zzk.cnblogs.com/s/blogpost'
    search_str = 'python 爬虫'
    for i in range(1, 3):
        html = request_html_by_query(url, Keywords=search_str,pageindex=i)
        #  每页下载5个blog
        search_arr = get_search_arr(html, 5)
        # 每个blog保存为word
        save_blogs_by_one(search_arr, i)



